"""Extract Markdown-shaped text from uploaded documents.

Chunking keys on headings, so every parser's job is to recover the document's
structure — not just its words. A PDF that arrives as one wall of text chunks
into one enormous passage and retrieves badly, which is why heading recovery
matters more here than raw extraction fidelity.
"""

import io
import logging
import re

from pypdf import PdfReader
from pypdf.errors import PdfReadError

logger = logging.getLogger(__name__)

MESSAGES = {
    "unsupported": "Formati i skedarit nuk mbeshtetet. Perdor PDF, DOCX, TXT ose MD.",
    "too_large": "Skedari eshte shume i madh. Kufiri eshte 10 MB.",
    "empty": "Skedari eshte bosh.",
    "unreadable": "Skedari nuk u lexua dot. Kontrollo qe nuk eshte i demtuar.",
    "no_text": "Nuk u gjet tekst ne skedar. Nese eshte PDF i skanuar, duhet OCR.",
    "no_headings": (
        "Dokumenti nuk ka tituj seksionesh. Shtoji tituj (p.sh. rreshta qe "
        "fillojne me ## ) qe permbajtja te ndahet si duhet."
    ),
    "no_chunks": (
        "Dokumenti nuk prodhoi asnje seksion te perdorshem. Kontrollo qe "
        "titujt kane tekst nen ta."
    ),
}

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

MAX_UPLOAD_BYTES = 10 * 1024 * 1024


class DocumentError(Exception):
    """A failure with a message the administrator can act on."""

    def __init__(self, message: str) -> None:
        super().__init__(message)
        self.message = message


def extension_of(filename: str) -> str:
    _, _, suffix = filename.rpartition(".")
    return f".{suffix.lower()}" if suffix else ""


def _read_pdf(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except (PdfReadError, OSError, ValueError):
        logger.exception("Could not read an uploaded PDF")
        raise DocumentError(MESSAGES["unreadable"]) from None

    return "\n\n".join(page.strip() for page in pages if page.strip())


def _read_docx(data: bytes) -> str:
    """Read a DOCX, mapping Word heading styles onto Markdown headings.

    Word carries structure in paragraph styles rather than in the text, so a
    naive text dump loses every heading and leaves the chunker with nothing to
    split on.
    """
    try:
        from docx import Document  # imported lazily: only uploads need it

        document = Document(io.BytesIO(data))
    except Exception:
        logger.exception("Could not read an uploaded DOCX")
        raise DocumentError(MESSAGES["unreadable"]) from None

    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style = (paragraph.style.name or "").lower() if paragraph.style else ""
        match = re.match(r"heading (\d)", style)

        if match:
            # Word's Heading 1 becomes Markdown `##`, not `#`. A document with
            # several `#` lines reads as several documents to the chunker,
            # which treats `#` as the title and resets its section stack at
            # each one — the result is zero chunks.
            level = min(int(match.group(1)) + 1, 6)
            lines.append(f"{'#' * level} {text}")
        elif style == "title":
            lines.append(f"# {text}")
        else:
            lines.append(text)

    # Tables hold most of the numbers in a policy document; losing them would
    # silently drop the facts an answer needs to cite.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append("| " + " | ".join(cells) + " |")

    return "\n\n".join(lines)


def _read_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise DocumentError(MESSAGES["unreadable"])


def _looks_like_heading(line: str) -> bool:
    """A short line that does not read as prose."""
    if not (0 < len(line) <= 80):
        return False

    if line.endswith((".", ",", ";", ":", "!", "?")):
        return False

    numbered = bool(re.match(r"^\d+(\.\d+)*[.)]?\s+\S", line))
    return numbered or line.istitle() or line.isupper()


def _promote_headings(text: str) -> str:
    """Give a plain-text document headings the chunker can split on.

    PDF and TXT carry no styling, so section titles have to be recovered from
    shape: short lines that do not end like a sentence, either numbered or
    capitalised. Blank lines are deliberately not required — extracted PDF
    text often has none at all, every line separated by a single newline.

    The first line becomes the document title, since the chunker needs one `#`
    to name the document and treats every later `#` as a new document.
    """
    if re.search(r"^#{1,6}\s+\S", text, flags=re.MULTILINE):
        return text  # already has Markdown headings

    lines = text.splitlines()
    promoted: list[str] = []
    title_taken = False

    for line in lines:
        stripped = line.strip()

        if not stripped:
            promoted.append(line)
            continue

        if not title_taken:
            promoted.append(f"# {stripped}")
            title_taken = True
            continue

        promoted.append(f"## {stripped}" if _looks_like_heading(stripped) else line)

    return "\n".join(promoted)


def parse_document(filename: str, data: bytes) -> str:
    """Turn an uploaded file into Markdown-shaped text ready for chunking."""
    extension = extension_of(filename)

    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentError(MESSAGES["unsupported"])

    if not data:
        raise DocumentError(MESSAGES["empty"])

    if len(data) > MAX_UPLOAD_BYTES:
        raise DocumentError(MESSAGES["too_large"])

    if extension == ".pdf":
        text = _promote_headings(_read_pdf(data))
    elif extension == ".docx":
        text = _read_docx(data)
    elif extension == ".txt":
        text = _promote_headings(_read_text(data))
    else:
        text = _read_text(data)

    if not text.strip():
        raise DocumentError(MESSAGES["no_text"])

    if not re.search(r"^#{1,6}\s+\S", text, flags=re.MULTILINE):
        # Better to reject than to index one unsearchable blob and let the
        # administrator discover it through bad answers later.
        raise DocumentError(MESSAGES["no_headings"])

    return text
