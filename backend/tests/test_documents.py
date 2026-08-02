import io

import pytest
from docx import Document

from app.services.documents import (
    MAX_UPLOAD_BYTES,
    DocumentError,
    _looks_like_heading,
    _promote_headings,
    parse_document,
)

MARKDOWN = "# Titulli\n\n## 1. Seksioni\n\nTekst i seksionit.\n"


def docx_bytes() -> bytes:
    """A Word document using real heading styles."""
    document = Document()
    document.add_heading("Politika Prove", level=0)
    document.add_heading("1. Seksioni i pare", level=1)
    document.add_paragraph("Teksti i seksionit te pare.")
    document.add_heading("1.1 Nenseksioni", level=2)
    document.add_paragraph("Teksti i nenseksionit.")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --- format and size validation ---


def test_an_unsupported_extension_is_rejected():
    with pytest.raises(DocumentError, match="Formati"):
        parse_document("skedar.pptx", b"whatever")


def test_a_file_without_an_extension_is_rejected():
    with pytest.raises(DocumentError, match="Formati"):
        parse_document("skedar", b"whatever")


def test_an_empty_file_is_rejected():
    with pytest.raises(DocumentError, match="bosh"):
        parse_document("prove.md", b"")


def test_an_oversized_file_is_rejected():
    with pytest.raises(DocumentError, match="madh"):
        parse_document("prove.md", b"x" * (MAX_UPLOAD_BYTES + 1))


def test_a_document_without_headings_is_rejected():
    with pytest.raises(DocumentError, match="tituj"):
        parse_document("prove.md", b"vetem tekst i rrjedhshem, pa asnje titull ketu.")


# --- Markdown passes through ---


def test_markdown_is_returned_unchanged():
    assert parse_document("prove.md", MARKDOWN.encode()) == MARKDOWN


def test_a_byte_order_mark_does_not_break_decoding():
    assert "Titulli" in parse_document("prove.md", MARKDOWN.encode("utf-8-sig"))


# --- DOCX ---


def test_word_heading_one_becomes_a_markdown_section():
    """Regression: mapping Heading 1 to `#` produced zero chunks.

    The chunker treats `#` as the document title and resets its section stack
    at each one, so a document with several of them has no sections at all.
    """
    text = parse_document("prove.docx", docx_bytes())

    assert "# Politika Prove" in text
    assert "## 1. Seksioni i pare" in text
    assert "### 1.1 Nenseksioni" in text
    # Exactly one line may start with a single `#`.
    assert sum(1 for line in text.splitlines() if line.startswith("# ")) == 1


def test_a_docx_chunks_into_sections():
    from app.services.chunking import chunk_markdown

    chunks = chunk_markdown(parse_document("prove.docx", docx_bytes())).chunks

    assert [c.heading_path[-1] for c in chunks] == ["1. Seksioni i pare", "1.1 Nenseksioni"]


def test_a_corrupt_docx_is_reported_as_unreadable():
    with pytest.raises(DocumentError, match="lexua"):
        parse_document("prove.docx", b"not a zip archive at all")


# --- plain text heading recovery ---


@pytest.mark.parametrize(
    "line",
    ["1. Orari i kantines", "2.1 Nenseksioni", "TITULL I MADH", "Politika E Kantines"],
)
def test_heading_shapes_are_recognised(line: str):
    assert _looks_like_heading(line)


@pytest.mark.parametrize(
    "line",
    [
        "Kantina eshte e hapur nga ora 11:30 deri ne 14:30.",
        "kjo eshte fjali e zakonshme",
        "Nje rresht qe mbaron me dy pika:",
        "x" * 120,
    ],
)
def test_prose_is_not_mistaken_for_a_heading(line: str):
    assert not _looks_like_heading(line)


def test_headings_are_recovered_without_blank_lines():
    """Regression: extracted PDF text often has no blank lines at all."""
    text = "Politika e Kantines\n1. Orari\nKantina eshte e hapur nga ora 11:30.\n2. Cmimet\nVakti kushton 300 leke."

    promoted = _promote_headings(text)

    assert promoted.splitlines()[0] == "# Politika e Kantines"
    assert "## 1. Orari" in promoted
    assert "## 2. Cmimet" in promoted


def test_only_the_first_line_becomes_the_title():
    promoted = _promote_headings("Titulli\n1. Nje\nTekst.\n2. Dy\nTekst.")

    assert sum(1 for line in promoted.splitlines() if line.startswith("# ")) == 1


def test_existing_markdown_headings_are_left_alone():
    assert _promote_headings(MARKDOWN) == MARKDOWN


def test_a_text_file_chunks_after_heading_recovery():
    from app.services.chunking import chunk_markdown

    raw = "Politika Prove\n1. Seksioni\nTekst i seksionit te pare ketu.\n"

    chunks = chunk_markdown(parse_document("prove.txt", raw.encode())).chunks

    assert [c.heading_path[-1] for c in chunks] == ["1. Seksioni"]
